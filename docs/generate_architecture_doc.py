"""生成软件体系结构文档（Word 格式）

运行方式：
    python docs/generate_architecture_doc.py

生成文件：
    docs/software_architecture.docx
"""

import os
import sys

# 确保项目根目录在 sys.path 中
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── 颜色常量 ────────────────────────────────────────────
BLUE_BG       = "D6E4F0"   # 应用层背景
BLUE_HEADER   = "4472C4"   # 应用层标题
BLUE_CELL     = "B4C6E7"   # 应用层单元格

ORANGE_BG     = "FBE5D6"   # 模型层背景
ORANGE_HEADER = "ED7D31"   # 模型层标题
ORANGE_CELL   = "F4B183"   # 模型层单元格

GREEN_BG      = "E2EFDA"   # 数据层背景
GREEN_HEADER  = "70AD47"   # 数据层标题
GREEN_CELL    = "A9D18E"   # 数据层单元格

GRAY_BG       = "F2F2F2"   # 箭头行背景
WHITE         = "FFFFFF"

FLOW_TRAIN    = "DAE3F3"   # 训练管线颜色
FLOW_INFER    = "E2EFDA"   # 推理管线颜色
FLOW_HUB      = "FFF2CC"   # 枢纽颜色

TABLE_BORDER  = "BFBFBF"
DARK_TEXT      = "333333"
HEADER_TEXT    = "FFFFFF"


# ─── 工具函数 ────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=10, color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name="微软雅黑"):
    """设置单元格文本及格式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # 垂直居中
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_border(table, color="BFBFBF"):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def remove_table_borders(table):
    """移除表格所有边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return heading


def add_body_text(doc, text, bold=False, size=11):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


# ─── 第一部分：系统架构图 ─────────────────────────────────

def build_architecture_diagram(doc):
    """生成三层架构图（用表格模拟）"""
    add_heading_styled(doc, "一、系统架构图", level=1)
    add_body_text(doc, "本系统采用经典的三层架构设计，自上而下分为应用层、模型层和数据层。各层之间通过明确的接口进行交互，层间依赖关系通过箭头标注。")
    doc.add_paragraph()

    # 架构图主体表格（7 行 1 列：应用层标题、应用层内容、箭头、模型层标题、模型层内容、箭头、数据层内容）
    table = doc.add_table(rows=7, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(16)

    # ── 应用层标题 ──
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_HEADER)
    set_cell_text(cell, "应用层（Application Layer）", bold=True, size=12, color=HEADER_TEXT)

    # ── 应用层内容（嵌套表格） ──
    cell = table.cell(1, 0)
    set_cell_shading(cell, BLUE_BG)
    inner = cell.add_table(rows=1, cols=2)
    remove_table_borders(inner)
    inner.autofit = True

    c1 = inner.cell(0, 0)
    set_cell_shading(c1, BLUE_CELL)
    set_cell_text(c1, "Gradio Web 界面\nsrc/app.py\n图片上传 · Top-3 结果展示", size=9, color=DARK_TEXT)

    c2 = inner.cell(0, 1)
    set_cell_shading(c2, BLUE_CELL)
    set_cell_text(c2, "训练入口\nrun_train.py\n数据分析 → 训练 → 评估流水线", size=9, color=DARK_TEXT)

    # ── 箭头行 1 ──
    cell = table.cell(2, 0)
    set_cell_shading(cell, GRAY_BG)
    set_cell_text(cell, "▼  调用模型层接口  ▼", size=10, color="666666")

    # ── 模型层标题 ──
    cell = table.cell(3, 0)
    set_cell_shading(cell, ORANGE_HEADER)
    set_cell_text(cell, "模型层（Model Layer）", bold=True, size=12, color=HEADER_TEXT)

    # ── 模型层内容 ──
    cell = table.cell(4, 0)
    set_cell_shading(cell, ORANGE_BG)
    inner = cell.add_table(rows=2, cols=3)
    remove_table_borders(inner)

    cells_data = [
        ("推理模块\nsrc/predict.py\n模型加载 · 单张推理", ORANGE_CELL),
        ("训练模块\nsrc/train.py\n训练循环 · 验证 · 保存", ORANGE_CELL),
        ("评估模块\nsrc/evaluate.py\n测试评估 · 混淆矩阵", ORANGE_CELL),
        ("模型定义\nsrc/model.py\nEfficientNet-B0 / ResNet-50", ORANGE_CELL),
        ("", WHITE),  # 空
        ("", WHITE),  # 空
    ]
    for i, (txt, clr) in enumerate(cells_data):
        r, c = divmod(i, 3)
        inner_cell = inner.cell(r, c)
        set_cell_shading(inner_cell, clr)
        if txt:
            set_cell_text(inner_cell, txt, size=9, color=DARK_TEXT)

    # ── 箭头行 2 ──
    cell = table.cell(5, 0)
    set_cell_shading(cell, GRAY_BG)
    set_cell_text(cell, "▼  读写数据层资源  ▼", size=10, color="666666")

    # ── 数据层标题+内容（合并） ──
    cell = table.cell(6, 0)
    set_cell_shading(cell, GREEN_BG)

    # 数据层标题
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("数据层（Data Layer）")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2E7D32")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 数据层内嵌表格
    inner = cell.add_table(rows=2, cols=4)
    remove_table_borders(inner)

    data_items = [
        ("数据集\narchive/\ntrain/valid/test\n100 类", GREEN_CELL),
        ("数据管道\nsrc/dataset.py\nDataLoader · 数据增强", GREEN_CELL),
        ("全局配置\nsrc/config.py\n路径常量 · 超参数", GREEN_CELL),
        ("持久化存储\noutputs/\nmodels/ + figures/", GREEN_CELL),
    ]
    for i, (txt, clr) in enumerate(data_items):
        r, c = divmod(i, 4)
        inner_cell = inner.cell(r, c)
        set_cell_shading(inner_cell, clr)
        set_cell_text(inner_cell, txt, size=9, color=DARK_TEXT)
    # 第二行空
    for c in range(4):
        set_cell_shading(inner.cell(1, c), WHITE)

    # ── 图例 ──
    doc.add_paragraph()
    legend_table = doc.add_table(rows=1, cols=3)
    legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(legend_table)

    legends = [
        ("■ 应用层", BLUE_HEADER, HEADER_TEXT),
        ("■ 模型层", ORANGE_HEADER, HEADER_TEXT),
        ("■ 数据层", GREEN_HEADER, HEADER_TEXT),
    ]
    for i, (txt, bg, fg) in enumerate(legends):
        cell = legend_table.cell(0, i)
        set_cell_shading(cell, bg)
        set_cell_text(cell, txt, bold=True, size=10, color=fg)


# ─── 第二部分：模块划分说明 ───────────────────────────────

MODULE_INFO = [
    {
        "name": "配置模块",
        "path": "src/config.py",
        "desc": "集中管理项目所有路径常量与超参数，供其他模块统一引用",
        "interfaces": [
            ("PROJECT_ROOT", "—", "Path", "项目根目录"),
            ("DATA_DIR", "—", "Path", "数据集根路径 (archive/)"),
            ("TRAIN_DIR / VALID_DIR / TEST_DIR", "—", "Path", "训练/验证/测试集路径"),
            ("OUTPUT_DIR / MODEL_DIR / FIGURE_DIR", "—", "Path", "输出目录路径"),
            ("NUM_CLASSES", "—", "int = 100", "分类类别数"),
            ("BATCH_SIZE", "—", "int = 32", "批大小"),
            ("NUM_EPOCHS", "—", "int = 15", "训练轮数"),
            ("LEARNING_RATE", "—", "float = 1e-4", "初始学习率"),
            ("IMAGE_SIZE", "—", "int = 224", "输入图像尺寸"),
            ("BEST_MODEL_PATH", "—", "Path", "最佳模型保存路径"),
        ],
        "deps": "pathlib",
    },
    {
        "name": "数据管道模块",
        "path": "src/dataset.py",
        "desc": "负责数据加载、预处理、数据增强以及数据集分析与可视化",
        "interfaces": [
            ("get_train_transforms", "—", "transforms.Compose", "训练集数据增强流水线（随机裁剪/翻转/旋转/色彩抖动）"),
            ("get_val_transforms", "—", "transforms.Compose", "验证/测试集预处理流水线（Resize + CenterCrop + 归一化）"),
            ("get_dataloaders", "batch_size, num_workers", "tuple[DataLoader, DataLoader, DataLoader, list]", "创建训练/验证/测试 DataLoader 及类别名"),
            ("analyze_dataset", "—", "None", "统计数据集类别分布并保存柱状图"),
            ("visualize_samples", "num_per_class, max_classes", "None", "生成训练集样本网格展示图"),
        ],
        "deps": "config, PIL, torch, torchvision, matplotlib, numpy",
    },
    {
        "name": "模型定义模块",
        "path": "src/model.py",
        "desc": "定义 EfficientNet-B0 和 ResNet-50 两种模型架构及其自定义分类头",
        "interfaces": [
            ("create_efficientnet_b0", "num_classes=100, pretrained=True", "nn.Module", "创建 EfficientNet-B0 模型（自定义分类头：Dropout→Linear→ReLU→Dropout→Linear）"),
            ("create_resnet50", "num_classes=100, pretrained=True, freeze_backbone=True", "nn.Module", "创建 ResNet-50 模型（可选冻结骨干网络）"),
            ("count_parameters", "model", "tuple[int, int]", "统计模型总参数量和可训练参数量"),
        ],
        "deps": "config, torch, torchvision",
    },
    {
        "name": "训练模块",
        "path": "src/train.py",
        "desc": "实现完整的训练与验证循环，包括单轮训练、验证评估和完整训练流水线",
        "interfaces": [
            ("train_one_epoch", "model, loader, criterion, optimizer, device", "tuple[float, float]", "单轮训练，返回 (平均损失, 准确率%)"),
            ("validate", "model, loader, criterion, device", "tuple[float, float]", "验证集评估，返回 (平均损失, 准确率%)"),
            ("train_model", "—", "dict", "完整训练流水线，返回训练历史字典"),
        ],
        "deps": "config, dataset, model, torch, tqdm",
    },
    {
        "name": "评估模块",
        "path": "src/evaluate.py",
        "desc": "提供训练曲线绘制、测试集评估、混淆矩阵生成和错误分析功能",
        "interfaces": [
            ("plot_training_history", "history, save_path=None", "None", "绘制 Loss/Acc/LR 训练曲线并保存"),
            ("evaluate_on_test", "—", "tuple[list, list, list, list]", "测试集完整评估，返回标签/预测/概率/类别名"),
            ("plot_confusion_matrix", "all_labels, all_preds, class_names", "None", "绘制混淆矩阵和每类准确率柱状图"),
            ("analyze_errors", "all_labels, all_preds, class_names, top_k=10", "list[tuple]", "分析最易混淆的类别对"),
            ("visualize_errors", "all_labels, all_preds, class_names, top_pairs=5", "None", "可视化错误分类样本图片"),
        ],
        "deps": "config, dataset, model, torch, sklearn, matplotlib, seaborn, PIL",
    },
    {
        "name": "推理模块",
        "path": "src/predict.py",
        "desc": "封装模型加载与单张图片推理的完整流程",
        "interfaces": [
            ("load_model", "model_path=None", "tuple[nn.Module, list, device]", "从 checkpoint 加载模型、类别名和设备"),
            ("predict_image", "model, image, class_names, device, top_k=3", "list[dict]", "单张图片推理，返回 Top-K 分类结果"),
        ],
        "deps": "config, dataset, model, torch, PIL",
    },
    {
        "name": "Web 应用模块",
        "path": "src/app.py",
        "desc": "基于 Gradio 构建 Web 交互界面，提供图片上传和实时分类功能",
        "interfaces": [
            ("_get_model", "—", "tuple[nn.Module, list, device]", "懒加载模型单例（首次调用时加载）"),
            ("predict", "image", "dict", "Gradio 预测回调，返回 {类别: 置信度}"),
            ("create_interface", "—", "tuple[Blocks, str]", "构建 Gradio 界面和 CSS 样式"),
        ],
        "deps": "config, predict, gradio, PIL",
    },
]


def build_module_descriptions(doc):
    """生成模块划分说明"""
    add_heading_styled(doc, "二、模块划分说明", level=1)
    add_body_text(doc, "系统共划分为 7 个功能模块，每个模块职责单一、接口明确。以下逐一列出各模块的职责描述、对外接口和依赖关系。")
    doc.add_paragraph()

    for mod in MODULE_INFO:
        # 模块标题
        add_heading_styled(doc, f"{mod['name']}（{mod['path']}）", level=2)

        # 职责描述
        p = doc.add_paragraph()
        run = p.add_run("职责：")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run2 = p.add_run(mod["desc"])
        run2.font.size = Pt(11)
        run2.font.name = "微软雅黑"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 接口表
        p = doc.add_paragraph()
        run = p.add_run("对外接口：")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        iface_table = doc.add_table(rows=len(mod["interfaces"]) + 1, cols=4)
        iface_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_border(iface_table)

        # 表头
        headers = ["函数/变量", "输入参数", "返回值", "功能说明"]
        for j, h in enumerate(headers):
            cell = iface_table.cell(0, j)
            set_cell_shading(cell, "4472C4")
            set_cell_text(cell, h, bold=True, size=9, color=HEADER_TEXT)

        # 数据行
        for i, (fn, params, ret, desc) in enumerate(mod["interfaces"], start=1):
            for j, val in enumerate([fn, params, ret, desc]):
                cell = iface_table.cell(i, j)
                bg = WHITE if i % 2 == 0 else "F2F2F2"
                set_cell_shading(cell, bg)
                alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 3 else WD_ALIGN_PARAGRAPH.LEFT
                set_cell_text(cell, val, size=9, color=DARK_TEXT, alignment=alignment)

        # 依赖关系
        p = doc.add_paragraph()
        run = p.add_run("依赖：")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run2 = p.add_run(mod["deps"])
        run2.font.size = Pt(11)
        run2.font.name = "微软雅黑"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        doc.add_paragraph()  # 模块间空行


def build_dependency_diagram(doc):
    """模块依赖关系图"""
    add_heading_styled(doc, "模块依赖关系图", level=2)
    add_body_text(doc, "以下用文本箭头表示模块间的导入依赖关系（A → B 表示 A 导入 B）：")
    doc.add_paragraph()

    dep_lines = [
        "app.py ──→ predict.py ──→ model.py ──→ config.py",
        "  │               └──→ dataset.py ──→ config.py",
        "  └──→ config.py",
        "",
        "run_train.py ──→ train.py ──→ model.py ──→ config.py",
        "     │              └──→ dataset.py ──→ config.py",
        "     ├──→ evaluate.py ──→ model.py",
        "     │              └──→ dataset.py",
        "     └──→ dataset.py",
        "",
        "train.py ──→ evaluate.py ──→ model.py",
        "                            └──→ dataset.py",
    ]

    for line in dep_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)


# ─── 第三部分：数据流向图 ────────────────────────────────

def build_data_flow_diagram(doc):
    """生成数据流向图"""
    add_heading_styled(doc, "三、数据流向图", level=1)
    add_body_text(doc, "系统包含两条核心数据管线：训练管线和推理管线。两条管线通过 best_model.pth 文件作为连接枢纽，训练管线产出模型文件，推理管线消费模型文件。")
    doc.add_paragraph()

    # ── 训练管线 ──
    add_heading_styled(doc, "训练管线数据流", level=2)

    train_steps = [
        ("原始图片\narchive/train/", FLOW_TRAIN),
        ("数据增强\nRandomResizedCrop\nRandomHorizontalFlip\nRandomRotation\nColorJitter", FLOW_TRAIN),
        ("DataLoader\n批次加载\nshuffle=True", FLOW_TRAIN),
        ("前向传播\nEfficientNet-B0\n输出 logits", FLOW_TRAIN),
        ("计算损失\nCrossEntropyLoss", FLOW_TRAIN),
        ("反向传播\nloss.backward()", FLOW_TRAIN),
        ("权重更新\nAdam 优化器\nStepLR 调度", FLOW_TRAIN),
    ]

    train_table = doc.add_table(rows=1, cols=len(train_steps) * 2 - 1)
    train_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(train_table)

    for i, (txt, clr) in enumerate(train_steps):
        col = i * 2
        cell = train_table.cell(0, col)
        set_cell_shading(cell, clr)
        set_cell_text(cell, txt, size=8, color=DARK_TEXT)
        cell.width = Cm(2.2)
        # 箭头列
        if i < len(train_steps) - 1:
            arrow_cell = train_table.cell(0, col + 1)
            set_cell_text(arrow_cell, "→", size=12, color="999999")
            arrow_cell.width = Cm(0.5)

    doc.add_paragraph()

    # 训练管线后半段
    train_steps_2 = [
        ("验证评估\nvalidate()", FLOW_TRAIN),
        ("保存最佳模型\nbest_model.pth", FLOW_HUB),
        ("测试集评估\nevaluate_on_test()", FLOW_TRAIN),
        ("可视化输出\n混淆矩阵 · 错误分析\n训练曲线", FLOW_TRAIN),
    ]

    train_table_2 = doc.add_table(rows=1, cols=len(train_steps_2) * 2 - 1)
    train_table_2.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(train_table_2)

    for i, (txt, clr) in enumerate(train_steps_2):
        col = i * 2
        cell = train_table_2.cell(0, col)
        set_cell_shading(cell, clr)
        set_cell_text(cell, txt, size=8, color=DARK_TEXT)
        cell.width = Cm(3)
        if i < len(train_steps_2) - 1:
            arrow_cell = train_table_2.cell(0, col + 1)
            set_cell_text(arrow_cell, "→", size=12, color="999999")
            arrow_cell.width = Cm(0.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── 推理管线 ──
    add_heading_styled(doc, "推理管线数据流", level=2)

    infer_steps = [
        ("用户上传图片\nGradio Image 组件", FLOW_INFER),
        ("Gradio 接收\n统一转为 PIL Image", FLOW_INFER),
        ("图像预处理\nRGB 转换\nResize(256)\nCenterCrop(224)\nImageNet 归一化", FLOW_INFER),
        ("模型前向推理\nEfficientNet-B0\nno_grad()", FLOW_INFER),
        ("Softmax\n概率归一化", FLOW_INFER),
        ("Top-3 筛选\ntorch.topk(k=3)", FLOW_INFER),
        ("结果展示\nGradio Label 组件\n{类别: 置信度}", FLOW_INFER),
    ]

    infer_table = doc.add_table(rows=1, cols=len(infer_steps) * 2 - 1)
    infer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(infer_table)

    for i, (txt, clr) in enumerate(infer_steps):
        col = i * 2
        cell = infer_table.cell(0, col)
        set_cell_shading(cell, clr)
        set_cell_text(cell, txt, size=8, color=DARK_TEXT)
        cell.width = Cm(2.2)
        if i < len(infer_steps) - 1:
            arrow_cell = infer_table.cell(0, col + 1)
            set_cell_text(arrow_cell, "→", size=12, color="999999")
            arrow_cell.width = Cm(0.5)

    doc.add_paragraph()

    # ── 枢纽说明 ──
    hub_table = doc.add_table(rows=1, cols=1)
    hub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(hub_table)
    hub_cell = hub_table.cell(0, 0)
    set_cell_shading(hub_cell, FLOW_HUB)
    set_cell_text(
        hub_cell,
        "连接枢纽：best_model.pth\n训练管线产出  ──→  best_model.pth  ──→  推理管线消费",
        bold=True, size=11, color="7F6000",
    )

    doc.add_paragraph()

    # 图例
    legend_table = doc.add_table(rows=1, cols=3)
    legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(legend_table)

    legends = [
        ("■ 训练管线", "4472C4", HEADER_TEXT),
        ("■ 推理管线", "548235", HEADER_TEXT),
        ("■ 模型枢纽", "BF8F00", HEADER_TEXT),
    ]
    for i, (txt, bg, fg) in enumerate(legends):
        cell = legend_table.cell(0, i)
        set_cell_shading(cell, bg)
        set_cell_text(cell, txt, bold=True, size=10, color=fg)


# ─── 主函数 ──────────────────────────────────────────────

def main():
    doc = Document()

    # 页面设置：A4 纵向
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # 文档标题
    title = doc.add_heading("软件体系结构", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Sports Image Classification — 运动图像分类系统")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()  # 空行

    # 构建三部分内容
    build_architecture_diagram(doc)
    doc.add_page_break()

    build_module_descriptions(doc)
    build_dependency_diagram(doc)
    doc.add_page_break()

    build_data_flow_diagram(doc)

    # 保存
    output_path = os.path.join(os.path.dirname(__file__), "software_architecture.docx")
    doc.save(output_path)
    print(f"文档已生成: {output_path}")


if __name__ == "__main__":
    main()
